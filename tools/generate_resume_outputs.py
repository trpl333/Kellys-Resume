"""
Generate ATS-oriented DOCX + PDF resumes from data/kelly_resume_source.json.

Requirements: python-docx, reportlab (see pip install in project notes).
"""

from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "data" / "kelly_resume_source.json"
OUT_DIR = ROOT / "outputs"
RESUME_DIR = ROOT / "resume"
PUBLIC_DIR = ROOT / "public"

# Full + DOCX PDFs: fewer signature bullets for whitespace (JSON keeps full list for web/markdown).
RESUME_SIGNATURE_RENDER_LIMIT = 5

BODY_FONT_PT = 10.5
SECTION_HEADING_PT = 11
# Full Professional History PDF/DOCX layout (target 2-3 pages).
BODY_LINE_SPACING = 1.10
PARA_SPACE_AFTER_PT = 3
ROLE_SPACE_BEFORE_PT = 5
BULLET_SPACE_AFTER_PT = 3

# Full Professional History PDF (ReportLab) — spacing / rhythm aligned with Quick polish.
FULL_PAGE_MARGIN_INCH = 0.75
FULL_SUMMARY_SECTION_PRE_BREAK_PT = int(round(BODY_FONT_PT * BODY_LINE_SPACING * 1.25))
FULL_INTER_SECTION_SPACER_PT = 7
FULL_CERT_SECTION_PRE_BREAK_PT = 12
FULL_EXPERIENCE_PRE_BREAK_PT = 8
FULL_EDU_SECTION_PRE_BREAK_PT = 14
FULL_ROLE_SPACE_BEFORE_PT = 6

# Quick Resume (2-page) PDF only: slightly tighter vertical rhythm.
QUICK_BODY_LINE_SPACING = 1.06
QUICK_PARA_SPACE_AFTER_PT = 2
QUICK_BULLET_SPACE_AFTER_PT = 2
QUICK_ROLE_SPACE_BEFORE_PT = 4
QUICK_SECTION_HEADING_BEFORE_PT = 6
# Extra air after name/contact before first SUMMARY heading (Quick PDF only).
# ~1.25× quick body line (~11pt) so header reads separate from body; keeps 2-page target with h1 spaceBefore.
QUICK_SUMMARY_SECTION_PRE_BREAK_PT = int(round(BODY_FONT_PT * QUICK_BODY_LINE_SPACING * 1.25))
QUICK_PAGE_MARGIN_INCH = 0.72
# Horizontal gap between two columns in Quick Resume tables (points).
QUICK_TABLE_COL_GUTTER_PT = 14
# Extra air before Certifications block (Quick PDF): ~1.35× body line (~11pt).
QUICK_CERT_SECTION_PRE_BREAK_PT = 15
# Same pre-break before Education (separates from Professional Experience).
QUICK_EDU_SECTION_PRE_BREAK_PT = 15
# Extra space below Education heading only (Quick PDF; Option A vs default h1).
QUICK_EDU_HEADING_SPACE_AFTER_PT = 6
# Vertical gap between certification rows (distinct entries; points).
QUICK_CERT_ROW_GAP_PT = 8
# Certifications column divider (Quick PDF only).
QUICK_CERT_LINE_W_PT = 0.35
QUICK_CERT_DIVIDER_PAD_PT = 5


def _core_competency_lines(items: list[str], *, max_lines: int = 2) -> list[str]:
    """Split competencies into at most two lines (no single long pipe-delimited row)."""
    if not items:
        return []
    if len(items) == 1:
        return [items[0]]
    mid = (len(items) + 1) // 2
    sep = " · "
    lines = [sep.join(items[:mid]), sep.join(items[mid:])]
    return lines[:max_lines]


def _quick_selected_bullets(role: dict[str, Any]) -> list[str]:
    """
    Editorial picks for the 2-page Quick Resume (older roles compressed).
    Orange USD + Chapman: full bullets. Others: capped per product spec.
    """
    bs = list(role.get("bullets") or [])
    emp = role.get("employer") or ""
    site = role.get("school_site") or ""
    if emp == "Orange Unified School District":
        return bs
    if emp == "Chapman University":
        return bs
    if emp == "Saddleback Valley Unified School District":
        if len(bs) > 3:
            return [bs[0], bs[3]]
        return bs[: min(2, len(bs))]
    if emp == "Capistrano Unified School District" and "Riley" in site:
        if len(bs) >= 3:
            return [bs[1], bs[2]]
        return bs[: min(2, len(bs))]
    if emp == "Mission Hills Christian School":
        if len(bs) > 4:
            return [bs[0], bs[4]]
        return bs[: min(2, len(bs))]
    if emp == "Chino Unified School District" and "Anna Borba" in site:
        if len(bs) > 10:
            return [bs[1], bs[10]]
        if len(bs) > 2:
            return [bs[1], bs[-1]]
        return bs[: min(2, len(bs))]
    if emp == "Chino Unified School District" and "Glenmeade" in site:
        return [bs[0]] if bs else []
    return bs


def _strip_markup_tags(text: str) -> str:
    return (
        text.replace("<b>", "")
        .replace("</b>", "")
        .replace("<B>", "")
        .replace("</B>", "")
    )


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("—", "-")
    )


def _text_for_output(text: str) -> str:
    """XML-safe text for ReportLab; never pass through literal <b> tags."""
    return _escape(_strip_markup_tags(text))


def load_data() -> dict[str, Any]:
    return json.loads(DATA_PATH.read_text(encoding="utf-8"))


def write_kelly_resume_markdown(data: dict[str, Any], path: Path, label: str) -> None:
    """Single-file resume for version control (Kellys-Resume source of truth)."""
    c = data["contact"]
    name = data["meta"]["candidate_name"]
    lines: list[str] = [
        f"# {name} — Resume ({label})",
        "",
        "> Source: `data/kelly_resume_source.json` in this repo. Facts only; student stories are privacy-safe.",
        "",
        "## Header",
        "",
        f"- **{name}**",
        f"- {c['city_state_zip']} | {c['phone']} | {c['email']}",
        "",
        "## Summary",
        "",
    ]
    for p in data["summary"]:
        lines.append(p)
        lines.append("")
    lines.extend(["## Signature Impact", ""])
    for b in data["signature_impact"]:
        lines.append(f"- {b}")
    lines.extend(["", "## Core Competencies", ""])
    for cc_line in _core_competency_lines(data["core_competencies"]):
        lines.append(cc_line)
    lines.extend(["", "## Certifications & Credentials", ""])
    for item in data["certifications"]:
        lines.append(f"- {item['name']} — {item['date']}")
    lines.extend(["", "## Professional Experience", ""])
    for role in data["experience"]:
        lines.append(f"### {role['title']}")
        lines.append("")
        sub = role["employer"]
        if role.get("school_site"):
            sub += f", {role['school_site']}"
        sub += f" · {role['location']} · {role['start']}–{role['end']}"
        if role.get("employment_note"):
            sub += f" · *{role['employment_note']}*"
        if role.get("date_note"):
            sub += f" · *{role['date_note']}*"
        lines.append(sub)
        lines.append("")
        for b in role["bullets"]:
            lines.append(f"- {b}")
        lines.append("")
    lines.extend(["## Education", ""])
    for edu in data["education"]:
        lines.append(f"- **{edu['degree']}**")
        lines.append(f"  - {edu['school']}, {edu['location']}")
        lines.append(f"  - {edu['date']}")
    lines.append("")
    refs = data.get("references_available") or []
    if refs:
        lines.extend(["## References (available upon request)", ""])
        for r in refs:
            lines.append(f"- **{r['name']}** — {r['title']}; {r.get('email', '')}; {r.get('phone', '')}")
        lines.append("")
    path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def set_doc_defaults(document: Document) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(BODY_FONT_PT)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = BODY_LINE_SPACING


def add_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(SECTION_HEADING_PT)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = BODY_LINE_SPACING


def add_bullet(document: Document, text: str) -> None:
    """Hanging bullet (custom) with whitespace between bullets."""
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.22)
    pf.first_line_indent = Inches(-0.14)
    pf.space_after = Pt(BULLET_SPACE_AFTER_PT)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = BODY_LINE_SPACING

    bullet_run = p.add_run("\u2022 ")
    bullet_run.font.name = "Calibri"
    bullet_run.font.size = Pt(BODY_FONT_PT)

    for part in _split_bold_segments(_strip_markup_tags(text)):
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.name = "Calibri"
            r.font.size = Pt(BODY_FONT_PT)
        else:
            r = p.add_run(part)
            r.font.name = "Calibri"
            r.font.size = Pt(BODY_FONT_PT)


def _split_bold_segments(text: str) -> list[str]:
    """Very small helper: supports **bold** segments."""
    parts: list[str] = []
    buf = ""
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            if buf:
                parts.append(buf)
                buf = ""
            j = text.find("**", i + 2)
            if j == -1:
                buf += text[i]
                i += 1
                continue
            parts.append(text[i : j + 2])
            i = j + 2
        else:
            buf += text[i]
            i += 1
    if buf:
        parts.append(buf)
    return parts


def build_header(document: Document, data: dict[str, Any]) -> None:
    c = data["contact"]
    name = data["meta"]["candidate_name"]
    p = document.add_paragraph()
    r = p.add_run(name.upper())
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    line2 = f"{c['city_state_zip']} | {c['phone']} | {c['email']}"
    p2 = document.add_paragraph(line2)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p2.paragraph_format.line_spacing = BODY_LINE_SPACING
    for run in p2.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(BODY_FONT_PT)


def _riley_volunteer_glance_line(data: dict[str, Any]) -> str:
    for role in data["experience"]:
        if role.get("school_site") == "Riley Elementary":
            return (
                f"{role['title']} | {role['employer']}, {role['school_site']} | "
                f"{role['location']} | {role['start']} – {role['end']}"
            )
    return (
        "Volunteer | Riley Elementary, Capistrano Unified School District | "
        "Capistrano Valley, CA | Jun 2013 – Mar 2016"
    )


def add_education_entry_docx(document: Document, edu: dict[str, Any], *, last: bool = False) -> None:
    """One degree block: bold degree, then school/location and date on separate lines (compact paragraph)."""
    gap_after = Pt(3 if last else 5)
    p = document.add_paragraph()
    r1 = p.add_run(_strip_markup_tags(edu["degree"]))
    r1.bold = True
    r1.font.name = "Calibri"
    r1.font.size = Pt(BODY_FONT_PT)
    p.add_run("\n")
    r2 = p.add_run(f"{edu['school']}, {edu['location']}")
    r2.font.name = "Calibri"
    r2.font.size = Pt(BODY_FONT_PT)
    p.add_run("\n")
    r3 = p.add_run(edu["date"])
    r3.font.name = "Calibri"
    r3.font.size = Pt(BODY_FONT_PT)
    p.paragraph_format.space_after = gap_after
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = BODY_LINE_SPACING


def _rl_education_degree_block_markup(edu: dict[str, Any]) -> str:
    """One ReportLab paragraph per degree; bold via mini-html (user text only is escaped)."""
    deg = _escape(_strip_markup_tags(edu["degree"]))
    loc = _escape(_strip_markup_tags(f"{edu['school']}, {edu['location']}"))
    dt = _escape(_strip_markup_tags(edu["date"]))
    return f"<b>{deg}</b><br/>{loc}<br/>{dt}"


def _quick_pdf_usable_width_pt() -> float:
    """Content width inside Quick Resume margins (points)."""
    return float(LETTER[0]) - 2 * QUICK_PAGE_MARGIN_INCH * 72


def _full_pdf_usable_width_pt() -> float:
    """Content width inside Full Professional History PDF margins (points)."""
    m = FULL_PAGE_MARGIN_INCH * inch
    return float(LETTER[0]) - 2 * m


def _quick_education_cell_markup(edu: dict[str, Any]) -> str:
    """Single-line education cell: bold degree | school, city | date (Quick PDF tables)."""
    deg = _escape(_strip_markup_tags(edu["degree"]))
    loc = _escape(_strip_markup_tags(f"{edu['school']}, {edu['location']}"))
    dt = _escape(_strip_markup_tags(edu["date"]))
    return f"<b>{deg}</b> | {loc} | {dt}"


def _quick_cert_cell_markup(name: str, date: str) -> str:
    """Bold credential title; date on second line, smaller and muted."""
    nm = _escape(_strip_markup_tags(name))
    dt = _escape(_strip_markup_tags(date))
    return (
        f"<b>{nm}</b><br/>"
        f'<font name="Helvetica" size="9" color="#6B6B6B">{dt}</font>'
    )


def build_resume_full_docx(data: dict[str, Any], path: Path) -> None:
    document = Document()
    set_doc_defaults(document)
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    build_header(document, data)

    add_heading(document, "Summary")
    for line in data["summary"]:
        p = document.add_paragraph(_strip_markup_tags(line))
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(BODY_FONT_PT)
        p.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = BODY_LINE_SPACING

    add_heading(document, "Signature Impact")
    for b in data["signature_impact"][:RESUME_SIGNATURE_RENDER_LIMIT]:
        add_bullet(document, b)

    add_heading(document, "Core Competencies")
    for cc_line in _core_competency_lines(data["core_competencies"]):
        pcc = document.add_paragraph(cc_line)
        for run in pcc.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(BODY_FONT_PT)
        pcc.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
        pcc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pcc.paragraph_format.line_spacing = BODY_LINE_SPACING

    add_heading(document, "Certifications & Credentials")
    for item in data["certifications"]:
        add_bullet(document, f"{item['name']} — {item['date']}")

    add_heading(document, "Professional Experience")
    for idx, role in enumerate(data["experience"]):
        header = _strip_markup_tags(
            f"{role['title']} | {role['employer']}"
            + (f", {role['school_site']}" if role.get("school_site") else "")
            + f" | {role['location']} | {role['start']} – {role['end']}"
            + (f" ({role['employment_note']})" if role.get("employment_note") else "")
            + (f" [{role['date_note']}]" if role.get("date_note") else "")
        )

        p = document.add_paragraph()
        if idx > 0:
            p.paragraph_format.space_before = Pt(ROLE_SPACE_BEFORE_PT)
        r = p.add_run(header)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(BODY_FONT_PT)
        p.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = BODY_LINE_SPACING

        for b in role["bullets"]:
            add_bullet(document, b)

    add_heading(document, "Education")
    edu_list = data["education"]
    for i, edu in enumerate(edu_list):
        add_education_entry_docx(document, edu, last=(i == len(edu_list) - 1))

    leadership = data.get("leadership_committees") or []
    if leadership:
        add_heading(document, "Selected Leadership & Committees")
        for item in leadership:
            add_bullet(document, item)

    document.save(path)


def build_executive_snapshot_docx(data: dict[str, Any], path: Path) -> None:
    document = Document()
    set_doc_defaults(document)
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    build_header(document, data)

    p = document.add_paragraph()
    r = p.add_run("Executive Snapshot")
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    add_heading(document, "Headline")
    headline = (
        "Special education teacher and Chapman University field supervisor with 26+ years across "
        "Special Day Class (SDC), Resource Specialist Program (RSP)/Inclusion, and general education (TK–8)."
    )
    p2 = document.add_paragraph(headline)
    for run in p2.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(BODY_FONT_PT)
    p2.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p2.paragraph_format.line_spacing = BODY_LINE_SPACING

    add_heading(document, "Signature Impact (selected)")
    for b in data["signature_impact"][:RESUME_SIGNATURE_RENDER_LIMIT]:
        add_bullet(document, b)

    add_heading(document, "Roles at a Glance")
    roles = [
        "Education Specialist (SDC + RSP/Inclusion) | Orange USD, Olive Elementary | Aug 2017 – Present",
        "Field Supervisor (SPED & Multiple Subject) | Chapman University | Aug 2020 – Present (caseload-based)",
        "SPED Instructional Aide | Saddleback Valley USD, RSM Intermediate | Mar 2016 – Jun 2017",
        _riley_volunteer_glance_line(data),
        "Junior High English Teacher | Mission Hills Christian School | Jun 2008 – Jun 2013",
        "Elementary Teacher | Chino USD, Anna Borba Elementary | Aug 1994 – Jun 2008",
        "Sixth Grade Teacher | Glenmeade Elementary, Chino Hills | Oct 1991 – Aug 1994",
    ]
    for line in roles:
        add_bullet(document, line)

    add_heading(document, "Credentials (quick view)")
    edu_quick = f"{data['education'][0]['degree']}, {data['education'][0]['school']} ({data['education'][0]['date']})"
    p3 = document.add_paragraph(edu_quick)
    for run in p3.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(BODY_FONT_PT)
    p3.paragraph_format.space_after = Pt(PARA_SPACE_AFTER_PT)
    p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p3.paragraph_format.line_spacing = BODY_LINE_SPACING
    for item in data["certifications"]:
        add_bullet(document, f"{item['name']} — {item['date']}")

    document.save(path)


def _rl_styles() -> dict[str, ParagraphStyle]:
    """ReportLab paragraph styles aligned to full resume DOCX (10.5pt body)."""
    base = getSampleStyleSheet()
    styles: dict[str, ParagraphStyle] = {}
    fs = float(BODY_FONT_PT)
    leading = round(fs * BODY_LINE_SPACING, 2)

    styles["name"] = ParagraphStyle(
        "name",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=16,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#111111"),
        spaceAfter=4,
    )
    styles["contact"] = ParagraphStyle(
        "contact",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=leading,
        alignment=TA_LEFT,
        spaceAfter=4,
    )
    styles["h1"] = ParagraphStyle(
        "h1",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=float(SECTION_HEADING_PT),
        leading=float(SECTION_HEADING_PT) + 1,
        textColor=colors.HexColor("#111111"),
        spaceBefore=8,
        spaceAfter=PARA_SPACE_AFTER_PT,
    )
    styles["body"] = ParagraphStyle(
        "body",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=leading,
        alignment=TA_LEFT,
        spaceAfter=PARA_SPACE_AFTER_PT,
    )
    styles["role_header"] = ParagraphStyle(
        "role_header",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=fs,
        leading=leading,
        alignment=TA_LEFT,
        spaceAfter=PARA_SPACE_AFTER_PT,
    )
    styles["edu_block"] = ParagraphStyle(
        "edu_block",
        parent=styles["body"],
        spaceAfter=4,
    )
    styles["edu_date_last"] = ParagraphStyle(
        "edu_date_last",
        parent=styles["body"],
        spaceAfter=PARA_SPACE_AFTER_PT,
    )
    styles["bullet"] = ParagraphStyle(
        "bullet",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=leading,
        leftIndent=16,
        bulletIndent=6,
        alignment=TA_LEFT,
        spaceAfter=BULLET_SPACE_AFTER_PT,
    )
    # Core Competencies: same body font, slightly tighter vertical rhythm (scan block).
    styles["body_core_full"] = ParagraphStyle(
        "body_core_full",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=leading,
        alignment=TA_LEFT,
        spaceAfter=2,
    )
    # Centered section heading (Full PDF certifications — matches Quick visual language).
    styles["h1_cert_full"] = ParagraphStyle(
        "h1_cert_full",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=float(SECTION_HEADING_PT),
        leading=float(SECTION_HEADING_PT) + 1,
        textColor=colors.HexColor("#111111"),
        alignment=TA_CENTER,
        spaceBefore=6,
        spaceAfter=3,
    )
    # Education heading: left-aligned, extra air below heading (Quick-style polish).
    styles["h1_edu_full"] = ParagraphStyle(
        "h1_edu_full",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=float(SECTION_HEADING_PT),
        leading=float(SECTION_HEADING_PT) + 1,
        textColor=colors.HexColor("#111111"),
        alignment=TA_LEFT,
        spaceBefore=6,
        spaceAfter=6,
    )
    # Two-column table cells (Full PDF): mirror Quick table rhythm with full body leading.
    styles["full_table_cell"] = ParagraphStyle(
        "full_table_cell",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=leading + 0.25,
        alignment=TA_LEFT,
        leftIndent=0,
        rightIndent=0,
        spaceBefore=0,
        spaceAfter=2,
    )
    styles["full_cert_table_cell"] = ParagraphStyle(
        "full_cert_table_cell",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=round(fs * 1.18, 2),
        alignment=TA_LEFT,
        leftIndent=0,
        rightIndent=0,
        spaceBefore=0,
        spaceAfter=0,
    )
    return styles


def _rl_styles_quick() -> dict[str, ParagraphStyle]:
    """Tighter vertical rhythm for the 2-page Quick Resume PDF only."""
    base = getSampleStyleSheet()
    styles: dict[str, ParagraphStyle] = {}
    fs = float(BODY_FONT_PT)
    leading = round(fs * QUICK_BODY_LINE_SPACING, 2)

    styles["name"] = ParagraphStyle(
        "name_quick",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=16,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#111111"),
        spaceAfter=3,
    )
    styles["contact"] = ParagraphStyle(
        "contact_quick",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=leading,
        alignment=TA_LEFT,
        spaceAfter=3,
    )
    styles["h1"] = ParagraphStyle(
        "h1_quick",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=float(SECTION_HEADING_PT),
        leading=float(SECTION_HEADING_PT) + 1,
        textColor=colors.HexColor("#111111"),
        spaceBefore=QUICK_SECTION_HEADING_BEFORE_PT,
        spaceAfter=QUICK_PARA_SPACE_AFTER_PT,
    )
    # Certifications block only: centered over the two-column table (Quick PDF).
    styles["h1_cert_quick"] = ParagraphStyle(
        "h1_cert_quick",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=float(SECTION_HEADING_PT),
        leading=float(SECTION_HEADING_PT) + 1,
        textColor=colors.HexColor("#111111"),
        alignment=TA_CENTER,
        spaceBefore=QUICK_SECTION_HEADING_BEFORE_PT,
        spaceAfter=QUICK_PARA_SPACE_AFTER_PT,
    )
    # Education block only: left-aligned; more air below heading than default h1 (Option A).
    styles["h1_edu_quick"] = ParagraphStyle(
        "h1_edu_quick",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=float(SECTION_HEADING_PT),
        leading=float(SECTION_HEADING_PT) + 1,
        textColor=colors.HexColor("#111111"),
        alignment=TA_LEFT,
        spaceBefore=QUICK_SECTION_HEADING_BEFORE_PT,
        spaceAfter=QUICK_EDU_HEADING_SPACE_AFTER_PT,
    )
    styles["body"] = ParagraphStyle(
        "body_quick",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=leading,
        alignment=TA_LEFT,
        spaceAfter=QUICK_PARA_SPACE_AFTER_PT,
    )
    styles["role_header"] = ParagraphStyle(
        "role_header_quick",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=fs,
        leading=leading,
        alignment=TA_LEFT,
        spaceAfter=QUICK_PARA_SPACE_AFTER_PT,
    )
    styles["edu_block"] = ParagraphStyle(
        "edu_block_quick",
        parent=styles["body"],
        spaceAfter=3,
    )
    styles["edu_date_last"] = ParagraphStyle(
        "edu_date_last_quick",
        parent=styles["body"],
        spaceAfter=QUICK_PARA_SPACE_AFTER_PT,
    )
    styles["bullet"] = ParagraphStyle(
        "bullet_quick",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=leading,
        leftIndent=16,
        bulletIndent=6,
        alignment=TA_LEFT,
        spaceAfter=QUICK_BULLET_SPACE_AFTER_PT,
    )
    # Table cells: left-aligned body, no hanging indent (Quick PDF only).
    styles["quick_table_cell"] = ParagraphStyle(
        "quick_table_cell",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=leading + 0.25,
        alignment=TA_LEFT,
        leftIndent=0,
        rightIndent=0,
        spaceBefore=0,
        spaceAfter=2,
    )
    # Certifications cells: two-line entries (title + date); slightly roomier leading.
    styles["quick_cert_table_cell"] = ParagraphStyle(
        "quick_cert_table_cell",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=fs,
        leading=round(fs * 1.18, 2),
        alignment=TA_LEFT,
        leftIndent=0,
        rightIndent=0,
        spaceBefore=0,
        spaceAfter=0,
    )
    return styles


def _rl_bullet_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(_text_for_output(text), style, bulletText="\u2022")


def _quick_two_column_table_style() -> TableStyle:
    """Minimal padding; no grid lines. Gutter between columns via asymmetric padding."""
    g = QUICK_TABLE_COL_GUTTER_PT
    return TableStyle(
        [
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (0, -1), g / 2),
            ("LEFTPADDING", (1, 0), (1, -1), g / 2),
            ("RIGHTPADDING", (1, 0), (1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]
    )


def _quick_two_column_table_style_certs() -> TableStyle:
    """Certifications table: light vertical divider; padding around divider; row gap."""
    pad = QUICK_CERT_DIVIDER_PAD_PT
    line = colors.HexColor("#D4D4D4")
    gap = QUICK_CERT_ROW_GAP_PT
    return TableStyle(
        [
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (0, -1), pad),
            ("LEFTPADDING", (1, 0), (1, -1), pad),
            ("RIGHTPADDING", (1, 0), (1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 1),
            ("BOTTOMPADDING", (0, 0), (1, 0), gap),
            ("BOTTOMPADDING", (0, 1), (-1, -1), 3),
            ("LINEAFTER", (0, 0), (0, -1), QUICK_CERT_LINE_W_PT, line),
        ]
    )


def _append_reportlab_certifications_two_column(
    story: list[Any],
    cell_style: ParagraphStyle,
    certifications: list[dict[str, Any]],
    *,
    usable_width_pt: float,
    gutter_pt: float = QUICK_TABLE_COL_GUTTER_PT,
    table_h_align: str = "CENTER",
) -> None:
    """
    Shared two-column certifications grid (Quick + Full PDFs).
    Order follows `data["certifications"]`; cells are left-aligned text, no bullets.
    """
    col_w = (usable_width_pt - gutter_pt) / 2
    rows: list[list[Any]] = []
    for i in range(0, len(certifications), 2):
        left = Paragraph(
            _quick_cert_cell_markup(certifications[i]["name"], certifications[i]["date"]),
            cell_style,
        )
        if i + 1 < len(certifications):
            right = Paragraph(
                _quick_cert_cell_markup(certifications[i + 1]["name"], certifications[i + 1]["date"]),
                cell_style,
            )
        else:
            right = Paragraph("", cell_style)
        rows.append([left, right])
    tbl = Table(rows, colWidths=[col_w, col_w], hAlign=table_h_align)
    tbl.setStyle(_quick_two_column_table_style_certs())
    story.append(tbl)


def _append_reportlab_education_two_column(
    story: list[Any],
    cell_style: ParagraphStyle,
    edu_rows: list[dict[str, Any]],
    *,
    usable_width_pt: float,
    gutter_pt: float = QUICK_TABLE_COL_GUTTER_PT,
    table_h_align: str = "LEFT",
) -> None:
    """Shared two-column education grid (Quick + Full PDFs); no vertical divider."""
    col_w = (usable_width_pt - gutter_pt) / 2
    rows: list[list[Any]] = []
    for i in range(0, len(edu_rows), 2):
        left = Paragraph(_quick_education_cell_markup(edu_rows[i]), cell_style)
        if i + 1 < len(edu_rows):
            right = Paragraph(_quick_education_cell_markup(edu_rows[i + 1]), cell_style)
        else:
            right = Paragraph("", cell_style)
        rows.append([left, right])
    tbl = Table(rows, colWidths=[col_w, col_w], hAlign=table_h_align)
    tbl.setStyle(_quick_two_column_table_style())
    story.append(tbl)


def _append_quick_certifications_two_column(
    story: list[Any], styles: dict[str, ParagraphStyle], certifications: list[dict[str, Any]]
) -> None:
    """Quick PDF: two-column certifications with Quick margins."""
    _append_reportlab_certifications_two_column(
        story,
        styles["quick_cert_table_cell"],
        certifications,
        usable_width_pt=_quick_pdf_usable_width_pt(),
        table_h_align="CENTER",
    )


def _append_quick_education_two_column(
    story: list[Any], styles: dict[str, ParagraphStyle], edu_rows: list[dict[str, Any]]
) -> None:
    """Quick PDF: two-column education with Quick margins."""
    _append_reportlab_education_two_column(
        story,
        styles["quick_table_cell"],
        edu_rows,
        usable_width_pt=_quick_pdf_usable_width_pt(),
        table_h_align="LEFT",
    )


def build_resume_full_pdf(data: dict[str, Any], path: Path) -> None:
    styles = _rl_styles()
    story: list[Any] = []
    c = data["contact"]
    name = data["meta"]["candidate_name"]

    story.append(Paragraph(_text_for_output(name.upper()), styles["name"]))
    story.append(
        Paragraph(
            _text_for_output(f"{c['city_state_zip']} | {c['phone']} | {c['email']}"),
            styles["contact"],
        )
    )

    story.append(Spacer(1, FULL_SUMMARY_SECTION_PRE_BREAK_PT))
    story.append(Paragraph(_escape("SUMMARY"), styles["h1"]))
    for line in data["summary"]:
        story.append(Paragraph(_text_for_output(line), styles["body"]))

    story.append(Spacer(1, FULL_INTER_SECTION_SPACER_PT))
    story.append(Paragraph(_escape("SIGNATURE IMPACT"), styles["h1"]))
    for b in data["signature_impact"][:RESUME_SIGNATURE_RENDER_LIMIT]:
        story.append(_rl_bullet_paragraph(b, styles["bullet"]))

    story.append(Spacer(1, FULL_INTER_SECTION_SPACER_PT))
    story.append(Paragraph(_escape("CORE COMPETENCIES"), styles["h1"]))
    for cc_line in _core_competency_lines(data["core_competencies"]):
        story.append(Paragraph(_text_for_output(cc_line), styles["body_core_full"]))

    story.append(Spacer(1, FULL_CERT_SECTION_PRE_BREAK_PT))
    story.append(Paragraph(_escape("CERTIFICATIONS & CREDENTIALS"), styles["h1_cert_full"]))
    story.append(Spacer(1, 2))
    _append_reportlab_certifications_two_column(
        story,
        styles["full_cert_table_cell"],
        data["certifications"],
        usable_width_pt=_full_pdf_usable_width_pt(),
        table_h_align="CENTER",
    )

    story.append(Spacer(1, FULL_EXPERIENCE_PRE_BREAK_PT))
    story.append(Paragraph(_escape("PROFESSIONAL EXPERIENCE"), styles["h1"]))
    story.append(Spacer(1, 4))
    for idx, role in enumerate(data["experience"]):
        if idx > 0:
            story.append(Spacer(1, FULL_ROLE_SPACE_BEFORE_PT))
        header = _strip_markup_tags(
            f"{role['title']} | {role['employer']}"
            + (f", {role['school_site']}" if role.get("school_site") else "")
            + f" | {role['location']} | {role['start']} – {role['end']}"
            + (f" ({role['employment_note']})" if role.get("employment_note") else "")
            + (f" [{role['date_note']}]" if role.get("date_note") else "")
        )
        story.append(Paragraph(_text_for_output(header), styles["role_header"]))
        for b in role["bullets"]:
            story.append(_rl_bullet_paragraph(b, styles["bullet"]))

    story.append(Spacer(1, FULL_EDU_SECTION_PRE_BREAK_PT))
    story.append(Paragraph(_escape("EDUCATION"), styles["h1_edu_full"]))
    _append_reportlab_education_two_column(
        story,
        styles["full_table_cell"],
        data["education"],
        usable_width_pt=_full_pdf_usable_width_pt(),
        table_h_align="LEFT",
    )

    leadership = data.get("leadership_committees") or []
    if leadership:
        story.append(Spacer(1, FULL_INTER_SECTION_SPACER_PT))
        story.append(Paragraph(_escape("SELECTED LEADERSHIP & COMMITTEES"), styles["h1"]))
        for item in leadership:
            story.append(_rl_bullet_paragraph(item, styles["bullet"]))

    story.append(Spacer(1, 0.03 * inch))

    margin = FULL_PAGE_MARGIN_INCH * inch
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
    )
    doc.build(story)


def build_resume_quick_pdf(data: dict[str, Any], path: Path) -> None:
    """
    Two-page Quick Resume: summary, certifications (dual Clear), signature impact,
    full Orange + Chapman roles, compressed older roles, education. No core competencies.
    """
    styles = _rl_styles_quick()
    story: list[Any] = []
    c = data["contact"]
    name = data["meta"]["candidate_name"]

    story.append(Paragraph(_text_for_output(name.upper()), styles["name"]))
    story.append(
        Paragraph(
            _text_for_output(f"{c['city_state_zip']} | {c['phone']} | {c['email']}"),
            styles["contact"],
        )
    )

    story.append(Spacer(1, QUICK_SUMMARY_SECTION_PRE_BREAK_PT))
    story.append(Paragraph(_escape("SUMMARY"), styles["h1"]))
    for line in data["summary"]:
        story.append(Paragraph(_text_for_output(line), styles["body"]))

    story.append(Spacer(1, QUICK_CERT_SECTION_PRE_BREAK_PT))
    story.append(Paragraph(_escape("CERTIFICATIONS & CREDENTIALS"), styles["h1_cert_quick"]))
    story.append(Spacer(1, 2))
    _append_quick_certifications_two_column(story, styles, data["certifications"])

    story.append(Paragraph(_escape("SIGNATURE IMPACT"), styles["h1"]))
    for b in data["signature_impact"][:RESUME_SIGNATURE_RENDER_LIMIT]:
        story.append(_rl_bullet_paragraph(b, styles["bullet"]))

    story.append(Paragraph(_escape("PROFESSIONAL EXPERIENCE"), styles["h1"]))
    for idx, role in enumerate(data["experience"]):
        if idx > 0:
            story.append(Spacer(1, QUICK_ROLE_SPACE_BEFORE_PT))
        header = _strip_markup_tags(
            f"{role['title']} | {role['employer']}"
            + (f", {role['school_site']}" if role.get("school_site") else "")
            + f" | {role['location']} | {role['start']} – {role['end']}"
            + (f" ({role['employment_note']})" if role.get("employment_note") else "")
            + (f" [{role['date_note']}]" if role.get("date_note") else "")
        )
        story.append(Paragraph(_text_for_output(header), styles["role_header"]))
        for b in _quick_selected_bullets(role):
            story.append(_rl_bullet_paragraph(b, styles["bullet"]))

    story.append(Spacer(1, QUICK_EDU_SECTION_PRE_BREAK_PT))
    story.append(Paragraph(_escape("EDUCATION"), styles["h1_edu_quick"]))
    _append_quick_education_two_column(story, styles, data["education"])

    story.append(Spacer(1, 0.02 * inch))

    m = QUICK_PAGE_MARGIN_INCH * inch
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=m,
        rightMargin=m,
        topMargin=m,
        bottomMargin=m,
    )
    doc.build(story)


def build_executive_snapshot_pdf(data: dict[str, Any], path: Path) -> None:
    styles = _rl_styles()
    story: list[Any] = []
    c = data["contact"]
    name = data["meta"]["candidate_name"]

    story.append(Paragraph(_text_for_output(name.upper()), styles["name"]))
    story.append(
        Paragraph(
            _text_for_output(f"{c['city_state_zip']} | {c['phone']} | {c['email']}"),
            styles["contact"],
        )
    )

    story.append(Paragraph(_escape("EXECUTIVE SNAPSHOT"), styles["h1"]))
    story.append(
        Paragraph(
            _text_for_output(
                "Special education teacher and Chapman University field supervisor with 26+ years across "
                "Special Day Class (SDC), Resource Specialist Program (RSP)/Inclusion, and general education (TK–8)."
            ),
            styles["body"],
        )
    )

    story.append(Paragraph(_escape("SIGNATURE IMPACT (SELECTED)"), styles["h1"]))
    for b in data["signature_impact"][:RESUME_SIGNATURE_RENDER_LIMIT]:
        story.append(_rl_bullet_paragraph(b, styles["bullet"]))

    story.append(Paragraph(_escape("ROLES AT A GLANCE"), styles["h1"]))
    roles = [
        "Education Specialist (SDC + RSP/Inclusion) | Orange USD, Olive Elementary | Aug 2017 – Present",
        "Field Supervisor (SPED & Multiple Subject) | Chapman University | Aug 2020 – Present (caseload-based)",
        "SPED Instructional Aide | Saddleback Valley USD, RSM Intermediate | Mar 2016 – Jun 2017",
        _riley_volunteer_glance_line(data),
        "Junior High English Teacher | Mission Hills Christian School | Jun 2008 – Jun 2013",
        "Elementary Teacher | Chino USD, Anna Borba Elementary | Aug 1994 – Jun 2008",
        "Sixth Grade Teacher | Glenmeade Elementary, Chino Hills | Oct 1991 – Aug 1994",
    ]
    for line in roles:
        story.append(_rl_bullet_paragraph(line, styles["bullet"]))

    story.append(Paragraph(_escape("CREDENTIALS (QUICK VIEW)"), styles["h1"]))
    story.append(
        Paragraph(
            _text_for_output(
                f"{data['education'][0]['degree']}, {data['education'][0]['school']} ({data['education'][0]['date']})"
            ),
            styles["body"],
        )
    )
    for item in data["certifications"]:
        story.append(_rl_bullet_paragraph(f"{item['name']} — {item['date']}", styles["bullet"]))

    margin = 0.75 * inch
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
    )
    doc.build(story)


def _assert_pdf_no_literal_bold_markup(path: Path, reader: Any) -> None:
    extracted = "".join((page.extract_text() or "") for page in reader.pages)
    bad_tags = ("<b>", "</b>", "<B>", "</B>", "&lt;b&gt;", "&lt;/b&gt;")
    if any(t in extracted for t in bad_tags):
        raise RuntimeError(f"Resume PDF at {path} appears to contain literal bold markup in extracted text.")


def _validate_quick_resume_pdf(path: Path) -> None:
    """Quick Resume must not exceed 2 pages."""
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError:
        from PyPDF2 import PdfReader  # type: ignore[import-not-found]

    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    if page_count > 2:
        raise RuntimeError(f"Quick Resume exceeded 2 pages (got {page_count}) at {path}.")
    if page_count < 2:
        print(f"WARNING: Quick Resume is only {page_count} page(s) at {path} (target is 2).")
    _assert_pdf_no_literal_bold_markup(path, reader)


def _validate_full_resume_pdf(path: Path) -> None:
    """Full Professional History: allow 2-3 pages; warn beyond 3."""
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError:
        from PyPDF2 import PdfReader  # type: ignore[import-not-found]

    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    if page_count > 3:
        raise RuntimeError(
            f"Full Professional History exceeded 3 pages (got {page_count}) at {path}."
        )
    if page_count == 3:
        print(
            f"WARNING: Full Professional History is 3 pages at {path} (allowed range is 2-3)."
        )
    if page_count < 2:
        raise RuntimeError(
            f"Full Professional History expected at least 2 pages, got {page_count} at {path}."
        )
    _assert_pdf_no_literal_bold_markup(path, reader)


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    RESUME_DIR.mkdir(exist_ok=True)
    PUBLIC_DIR.mkdir(exist_ok=True)
    data = load_data()

    write_kelly_resume_markdown(data, RESUME_DIR / "kelly_resume_v3.md", "v3")
    write_kelly_resume_markdown(data, RESUME_DIR / "kelly_resume_v2.md", "v2")

    build_resume_full_docx(data, OUT_DIR / "resume_full.docx")
    build_executive_snapshot_docx(data, OUT_DIR / "executive_snapshot.docx")

    build_resume_full_pdf(data, OUT_DIR / "resume_full.pdf")
    build_resume_quick_pdf(data, OUT_DIR / "resume_2page.pdf")
    build_executive_snapshot_pdf(data, OUT_DIR / "executive_snapshot.pdf")

    shutil.copyfile(OUT_DIR / "resume_full.pdf", PUBLIC_DIR / "resume_full.pdf")
    shutil.copyfile(OUT_DIR / "resume_2page.pdf", PUBLIC_DIR / "resume_2page.pdf")
    _validate_full_resume_pdf(PUBLIC_DIR / "resume_full.pdf")
    _validate_quick_resume_pdf(PUBLIC_DIR / "resume_2page.pdf")

    print("Wrote:")
    print(" -", RESUME_DIR / "kelly_resume_v2.md")
    print(" -", RESUME_DIR / "kelly_resume_v3.md")
    print(" -", PUBLIC_DIR / "resume_full.pdf")
    print(" -", PUBLIC_DIR / "resume_2page.pdf")
    for name in [
        "resume_full.docx",
        "executive_snapshot.docx",
        "resume_full.pdf",
        "resume_2page.pdf",
        "executive_snapshot.pdf",
    ]:
        print(" -", OUT_DIR / name)


if __name__ == "__main__":
    main()
